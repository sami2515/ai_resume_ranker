"""
Resume Ingestion (project documentation Section 5.1, Step 1)
--------------------------------------------------------------
Extracts raw text from .docx (python-docx) and .pdf (pdfplumber) resumes,
and does light section-boundary detection so downstream extraction can
scope its search (e.g. only look for degrees under an "Education" header).

Testing note (Section 8.3): this module is deliberately defensive --
corrupt or non-standard files must return a clear error, never raise an
uncaught exception into the API layer.
"""

from __future__ import annotations
import re
from dataclasses import dataclass, field
from pathlib import Path

import docx


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed. Carries a user-safe message."""


# Common resume section header variants -> canonical section name.
SECTION_HEADERS = {
    # --- Education ---
    "education": "education",
    "academic background": "education",
    "academic qualifications": "education",
    "academics": "education",
    "educational background": "education",
    "educational qualifications": "education",
    "qualifications": "education",
    "degrees": "education",
    "schooling": "education",

    # --- Experience ---
    "professional experience": "experience",
    "work experience": "experience",
    "experience": "experience",
    "employment history": "experience",
    "career history": "experience",
    "work history": "experience",
    "job history": "experience",
    "relevant experience": "experience",
    "professional background": "experience",
    "internship experience": "experience",
    "internships": "experience",
    "industry experience": "experience",

    # --- Skills ---
    "skills": "skills",
    "technical skills": "skills",
    "core competencies": "skills",
    "key skills": "skills",
    "areas of expertise": "skills",
    "expertise": "skills",
    "competencies": "skills",
    "it skills": "skills",
    "computer skills": "skills",
    "technologies": "skills",
    "tools and technologies": "skills",
    "tools & technologies": "skills",
    "programming languages": "skills",
    "languages and frameworks": "skills",

    # --- Certifications ---
    "certifications": "certifications",
    "certification": "certifications",
    "courses": "certifications",
    "training": "certifications",
    "trainings": "certifications",
    "trainings & courses": "certifications",
    "professional development": "certifications",
    "professional training": "certifications",
    "licenses & certifications": "certifications",
    "licenses and certifications": "certifications",
    "workshops": "certifications",
    "online courses": "certifications",

    # --- Summary ---
    "summary": "summary",
    "professional summary": "summary",
    "objective": "summary",
    "career objective": "summary",
    "professional objective": "summary",
    "career summary": "summary",
    "executive summary": "summary",
    "profile": "summary",
    "professional profile": "summary",
    "about me": "summary",
    "overview": "summary",

    # --- Projects ---
    "projects": "projects",
    "key projects": "projects",
    "project experience": "projects",
    "academic projects": "projects",
    "personal projects": "projects",
    "freelance projects": "projects",

    # --- Languages (spoken) ---
    "languages": "languages",
    "language proficiency": "languages",
    "spoken languages": "languages",

    # --- References ---
    "references": "references",
    "referees": "references",
}


@dataclass
class ParsedResume:
    filename: str
    raw_text: str
    sections: dict = field(default_factory=dict)  # section_name -> text block

    @property
    def is_empty(self) -> bool:
        return len(self.raw_text.strip()) < 20


def _looks_like_header(line: str) -> str | None:
    """Return the canonical section name if `line` looks like a section header."""
    clean = line.strip().strip(":").lower()
    if not clean or len(clean) > 40:
        return None
    if clean in SECTION_HEADERS:
        return SECTION_HEADERS[clean]
    return None


def _split_into_sections(lines: list[str]) -> dict:
    sections: dict[str, list[str]] = {}
    current = "header"  # anything before the first recognized header
    for line in lines:
        header = _looks_like_header(line)
        if header:
            current = header
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def parse_docx(path: str | Path) -> ParsedResume:
    path = Path(path)
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 - deliberately broad, see module docstring
        raise ResumeParseError(
            f"Could not open '{path.name}'. The file may be corrupt, "
            f"password-protected, or not a valid .docx file."
        ) from exc

    lines = [p.text for p in document.paragraphs if p.text.strip()]

    # Resumes sometimes put content in tables rather than paragraphs.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())

    raw_text = "\n".join(lines)
    resume = ParsedResume(filename=path.name, raw_text=raw_text, sections=_split_into_sections(lines))

    if resume.is_empty:
        raise ResumeParseError(f"'{path.name}' appears to contain no readable text.")
    return resume


def parse_pdf(path: str | Path) -> ParsedResume:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ResumeParseError("PDF support requires 'pdfplumber' (pip install pdfplumber).") from exc

    path = Path(path)
    try:
        with pdfplumber.open(str(path)) as pdf:
            text_chunks = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError(
            f"Could not open '{path.name}'. The file may be corrupt, scanned-image-only, "
            f"or password-protected."
        ) from exc

    raw_text = "\n".join(text_chunks)
    lines = [ln for ln in raw_text.splitlines() if ln.strip()]
    resume = ParsedResume(filename=path.name, raw_text=raw_text, sections=_split_into_sections(lines))

    if resume.is_empty:
        raise ResumeParseError(
            f"'{path.name}' appears to contain no extractable text "
            f"(it may be a scanned image without OCR)."
        )
    return resume


ALLOWED_EXTENSIONS = {".docx", ".pdf"}


def parse_resume(path: str | Path) -> ParsedResume:
    """Dispatch to the right parser based on file extension. Raises ResumeParseError."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".pdf":
        return parse_pdf(path)
    raise ResumeParseError(
        f"Unsupported file type '{ext}' for '{path.name}'. Only .docx and .pdf are supported."
    )
